import os
from typing import List
from docx import Document as DocxDocument
from langchain.docstore.document import Document


def load_docx_documents(folder_path: str) -> List[Document]:
    """
    Загружает все .docx-файлы из указанной папки,
    извлекает текст и возвращает список объектов LangChain Document
    """
    documents = []

    for filename in os.listdir(folder_path):
        if  filename.endswith(".docx") and not filename.startswith("~$"):
            path = os.path.join(folder_path, filename)
            text = extract_text_from_docx(path)
            doc = Document(page_content=text, metadata={"source": filename})
            documents.append(doc)

    return documents


def extract_text_from_docx(filepath: str) -> str:
    """
    Извлекает текст из docx файла.
    """
    doc = DocxDocument(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text.strip())
    return "\n".join(filter(None, full_text))


# Пример использования
if __name__ == "__main__":
    docs = load_docx_documents(r"D:\R-STYLE WORK\study_rag")
    print(f"Загружено документов: {len(docs)}")
    print(docs[0].page_content[:1000])